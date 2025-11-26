import fitz  # PyMuPDF
from docx import Document
from pathlib import Path
from typing import List, Dict, Any
import base64
from io import BytesIO
from PIL import Image

class DocumentProcessor:
    """Converts PDF/DOC to pages with text, images, and tables"""
    
    async def convert_to_pages(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Convert document to structured pages
        
        Args:
            file_path: Path to PDF or DOC file
            
        Returns:
            List of page dictionaries with text, images, and metadata
        """
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            return await self._process_pdf(file_path)
        elif file_ext in ['.doc', '.docx']:
            return await self._process_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    async def _process_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """
        Extract text, images, and tables from PDF
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            List of page data
        """
        pages = []
        
        try:
            doc = fitz.open(pdf_path)
            
            for page_num, page in enumerate(doc, start=1):
                # Extract text
                text = page.get_text("text")
                
                # Extract images
                images = []
                image_list = page.get_images(full=True)
                
                for img_index, img_info in enumerate(image_list):
                    xref = img_info[0]
                    try:
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        
                        # Convert to base64 for storage/transmission
                        image_base64 = base64.b64encode(image_bytes).decode('utf-8')
                        
                        images.append({
                            "image_id": f"page_{page_num}_img_{img_index}",
                            "format": base_image["ext"],
                            "data": image_base64,
                            "page": page_num
                        })
                    except Exception as e:
                        print(f"Error extracting image on page {page_num}: {e}")
                
                # Get page dimensions for table detection later
                page_rect = page.rect
                
                page_data = {
                    "page_number": page_num,
                    "text": text,
                    "images": images,
                    "dimensions": {
                        "width": page_rect.width,
                        "height": page_rect.height
                    },
                    "raw_page": page  # Keep reference for table detection
                }
                
                pages.append(page_data)
            
            doc.close()
            
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")
        
        return pages
    
    async def _process_docx(self, docx_path: str) -> List[Dict[str, Any]]:
        """
        Extract text and images from DOCX
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            List of page data (DOCX doesn't have pages, so we create logical pages)
        """
        pages = []
        
        try:
            doc = Document(docx_path)
            
            # Extract all text
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract images
            images = []
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                        
                        images.append({
                            "image_id": f"doc_img_{len(images)}",
                            "format": rel.target_ref.split('.')[-1],
                            "data": image_base64
                        })
                    except Exception as e:
                        print(f"Error extracting DOCX image: {e}")
            
            # Split into logical pages (every 50 paragraphs or tables)
            paragraphs_per_page = 50
            total_paragraphs = len(doc.paragraphs)
            
            for i in range(0, total_paragraphs, paragraphs_per_page):
                page_text = '\n'.join(full_text[i:i+paragraphs_per_page])
                
                page_data = {
                    "page_number": (i // paragraphs_per_page) + 1,
                    "text": page_text,
                    "images": images if i == 0 else [],  # Include images only on first page
                    "dimensions": {"width": 0, "height": 0},  # DOCX doesn't have fixed dimensions
                    "raw_page": None
                }
                
                pages.append(page_data)
            
        except Exception as e:
            raise Exception(f"DOCX processing error: {str(e)}")
        
        return pages